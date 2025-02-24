"""Document service module for handling file operations and caching."""

import asyncio
import hashlib
import os
import time
import uuid
from functools import lru_cache
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import io
import logging
from docx import Document as DocxDocument

import aiofiles
import magic
from fastapi import UploadFile
import PyPDF2

from app.core.config import settings
from app.core.exceptions import (
    DocumentNotFoundError,
    InvalidFileTypeError,
    FileSizeLimitError,
)
from app.models.schemas import Document

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class DocumentService:
    """Service for managing document operations with caching."""

    def __init__(self) -> None:
        """Initialize document service with caching."""
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
        self.content_cache: Dict[str, Tuple[bytes, float]] = {}
        self.cache_ttl = 300  # Cache TTL in seconds (5 minutes)
        self.path_cache: Dict[str, Tuple[Path, float]] = {}
        self.batch_size = 10  # Number of files to process in parallel
        self.timing_metrics: Dict[str, float] = {}

    def _record_timing(self, step: str, start_time: float) -> float:
        """Record timing for a step and return new start time."""
        elapsed = time.time() - start_time
        self.timing_metrics[step] = elapsed
        return time.time()

    def _get_timing_summary(self) -> str:
        """Generate timing summary."""
        total_time = sum(self.timing_metrics.values())
        summary = ["Document Processing Time Breakdown:"]
        for step, duration in self.timing_metrics.items():
            percentage = (duration / total_time) * 100
            summary.append(
                f"- {step}: {duration:.2f}s ({percentage:.1f}%)"
            )
        summary.append(f"Total Time: {total_time:.2f}s")
        return "\n".join(summary)

    async def save_document(self, file: UploadFile) -> str:
        """Save an uploaded document and return its ID."""
        self.timing_metrics.clear()

        if not file.filename:
            raise InvalidFileTypeError(set(settings.ALLOWED_EXTENSIONS))

        try:
            # Validate file extension
            validation_start = time.time()
            ext = file.filename.split('.')[-1].lower()
            if ext not in settings.ALLOWED_EXTENSIONS:
                raise InvalidFileTypeError(set(settings.ALLOWED_EXTENSIONS))

            # Read and validate content
            content = await file.read()
            await file.seek(0)

            if len(content) > settings.MAX_UPLOAD_SIZE:
                raise FileSizeLimitError(settings.MAX_UPLOAD_SIZE)

            # Validate MIME type
            mime = magic.Magic(mime=True)
            file_type = mime.from_buffer(content)
            if not self._is_valid_mime_type(file_type):
                raise InvalidFileTypeError(set(settings.ALLOWED_EXTENSIONS))
            
            self._record_timing("Validation", validation_start)

            # Generate ID and prepare for save
            prep_start = time.time()
            document_id = str(uuid.uuid4())
            file_path = self.upload_dir / f"{document_id}.{ext}"
            content_hash = hashlib.sha256(content).hexdigest()
            self._record_timing("Preparation", prep_start)

            # Save file
            save_start = time.time()
            self.upload_dir.mkdir(parents=True, exist_ok=True)
            async with aiofiles.open(file_path, 'wb') as f:
                chunk_size = 1024 * 1024  # 1MB chunks
                while chunk := await file.read(chunk_size):
                    await f.write(chunk)
            self._record_timing("File Writing", save_start)

            # Verify and cache
            verify_start = time.time()
            if not await self._verify_file_hash(file_path, content_hash):
                await self._remove_file(file_path)
                raise ValueError("File verification failed")

            self._cache_content(document_id, content)
            self._cache_path(document_id, file_path)
            self._record_timing("Verification & Caching", verify_start)

            logger.debug(self._get_timing_summary())
            return document_id

        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            if 'file_path' in locals() and file_path.exists():
                await self._remove_file(file_path)
            raise

    async def get_document_path(self, document_id: str) -> Path:
        """Get the file path for a document ID."""
        cached_path = self._get_cached_path(document_id)
        if cached_path:
            return cached_path

        for ext in settings.ALLOWED_EXTENSIONS:
            path = self.upload_dir / f"{document_id}.{ext}"
            if path.exists():
                self._cache_path(document_id, path)
                return path
        raise DocumentNotFoundError(document_id)

    async def get_document_content(self, document_id: str) -> bytes:
        """Get document content with caching."""
        self.timing_metrics.clear()

        try:
            # Check cache
            cache_start = time.time()
            cached_content = self._get_cached_content(document_id)
            if cached_content:
                self._record_timing("Cache Retrieval", cache_start)
                return cached_content

            # Get file path
            path_start = time.time()
            path = await self.get_document_path(document_id)
            self._record_timing("Path Resolution", path_start)

            # Extract content
            extract_start = time.time()
            content = None
            extraction_method = "direct"

            # For text files, try direct reading first
            if path.suffix.lower() == '.txt':
                try:
                    async with aiofiles.open(path, 'r', encoding='utf-8') as f:
                        content = await f.read()
                except UnicodeDecodeError:
                    try:
                        async with aiofiles.open(path, 'rb') as f:
                            binary_content = await f.read()
                            content = binary_content.decode('latin-1')
                    except Exception as e:
                        logger.error(f"Error reading text file: {str(e)}")
                        extraction_method = "llm"

            elif path.suffix.lower() == '.pdf':
                try:
                    content = await self._extract_pdf_text(path)
                    if not content or len(content.strip()) < 100:
                        extraction_method = "llm"
                except Exception:
                    extraction_method = "llm"

            elif path.suffix.lower() == '.docx':
                try:
                    content = await self._extract_docx_text(path)
                    if not content or len(content.strip()) < 100:
                        extraction_method = "llm"
                except Exception:
                    extraction_method = "llm"

            # If direct extraction failed or content is insufficient, try LLM
            if extraction_method == "llm" or not content:
                content = await self._extract_with_llm(path)

            if not content:
                raise ValueError("Failed to extract content")

            timing_msg = (
                f"Content Extraction ({extraction_method})"
            )
            self._record_timing(timing_msg, extract_start)

            # Cache the content
            cache_start = time.time()
            content_bytes = content.encode() if isinstance(content, str) else content
            self._cache_content(document_id, content_bytes)
            self._record_timing("Content Caching", cache_start)

            logger.debug(self._get_timing_summary())
            return content_bytes

        except Exception as e:
            logger.error(f"Error getting document content: {str(e)}")
            raise

    async def _extract_pdf_text(self, path: Path) -> str:
        """Extract text content from a PDF file."""
        try:
            # Read the PDF file in binary mode
            logger.debug("Opening PDF file")
            async with aiofiles.open(path, 'rb') as file:
                content = await file.read()
                
            # Create a PDF reader object
            logger.debug("Creating PDF reader")
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            logger.debug(f"Number of pages: {len(pdf_reader.pages)}")
            
            # Extract text from all pages
            text_content = []
            for i, page in enumerate(pdf_reader.pages):
                logger.debug(f"Extracting text from page {i+1}")
                text = page.extract_text()
                if text:
                    text_content.append(text.strip())
                    logger.debug(f"Page {i+1} text length: {len(text)}")
            
            full_text = "\n\n".join(text_content)
            logger.debug(f"Total extracted text length: {len(full_text)}")
            return full_text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise ValueError(f"Error extracting PDF text: {str(e)}")

    async def _extract_docx_text(self, path: Path) -> str:
        """Extract text content from a DOCX file."""
        try:
            # Read the DOCX file
            logger.debug("Opening DOCX file")
            doc = DocxDocument(path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            full_text = "\n\n".join(paragraphs)
            logger.debug(f"Total extracted text length: {len(full_text)}")
            return full_text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise ValueError(f"Error extracting DOCX text: {str(e)}")

    async def _extract_with_llm(self, path: Path) -> str:
        """Extract text content using LLM for challenging documents."""
        try:
            # For text-based models, we'll try to read the file as text first
            try:
                async with aiofiles.open(path, 'r', encoding='utf-8') as f:
                    raw_content = await f.read()
            except UnicodeDecodeError:
                # If text reading fails, try reading as binary and decode with latin-1
                async with aiofiles.open(path, 'rb') as f:
                    content = await f.read()
                    raw_content = content.decode('latin-1', errors='ignore')

            # Create a specialized prompt for document extraction
            prompt = (
                "You are a document content extraction specialist. Your task is to:\n"
                "1. Extract and organize all text content from the provided document\n"
                "2. Preserve the logical structure and flow\n"
                "3. Properly format tables, lists, and other elements\n"
                "4. Include metadata like title and headers\n"
                "5. Remove any non-text elements or formatting artifacts\n\n"
                f"Document type: {path.suffix}\n"
                "Document content:\n\n"
                f"{raw_content[:10000]}"  # First 10K chars to stay within context limits
            )

            # Use a specialized model for document extraction
            from app.services.llm import LLMService
            llm_service = LLMService()
            
            # Use our default model (it's already initialized)
            extracted_content = await llm_service._get_completion(prompt)
            
            if not extracted_content:
                raise ValueError("LLM extraction returned empty content")
            
            return str(extracted_content)

        except Exception as e:
            logger.error(f"LLM extraction failed: {str(e)}")
            raise ValueError(f"LLM extraction failed: {str(e)}")

    async def list_documents(self) -> List[Document]:
        """List all uploaded documents with parallel processing."""
        files = [
            f for f in self.upload_dir.glob('*')
            if any(f.name.endswith(ext) for ext in settings.ALLOWED_EXTENSIONS)
        ]

        documents = []
        for i in range(0, len(files), self.batch_size):
            batch = files[i:i + self.batch_size]
            tasks = [self._process_document(file_path) for file_path in batch]
            batch_results = await asyncio.gather(*tasks)
            documents.extend(batch_results)

        return documents

    async def _process_document(self, file_path: Path) -> Document:
        """Process a single document file."""
        doc_id = file_path.stem
        stat = file_path.stat()
        return Document(
            id=doc_id,
            filename=file_path.name,
            size=stat.st_size,
            content_type=self._get_content_type(file_path.name)
        )

    def _cache_content(self, document_id: str, content: bytes) -> None:
        """Cache document content with timestamp."""
        self.content_cache[document_id] = (content, time.time())

    def _get_cached_content(self, document_id: str) -> Optional[bytes]:
        """Get cached content if not expired."""
        if document_id in self.content_cache:
            content, timestamp = self.content_cache[document_id]
            if time.time() - timestamp < self.cache_ttl:
                return content
            del self.content_cache[document_id]
        return None

    def _cache_path(self, document_id: str, path: Path) -> None:
        """Cache document path with timestamp."""
        self.path_cache[document_id] = (path, time.time())

    def _get_cached_path(self, document_id: str) -> Optional[Path]:
        """Get cached path if not expired."""
        if document_id in self.path_cache:
            path, timestamp = self.path_cache[document_id]
            if time.time() - timestamp < self.cache_ttl:
                return path
            del self.path_cache[document_id]
        return None

    @lru_cache(maxsize=128)
    def _get_content_type(self, filename: str) -> str:
        """Get MIME type for a filename (cached)."""
        ext = filename.split('.')[-1].lower()
        content_types = {
            'txt': 'text/plain',
            'pdf': 'application/pdf',
            'doc': 'application/msword',
            'docx': (
                'application/vnd.openxmlformats-officedocument.'
                'wordprocessingml.document'
            )
        }
        return content_types.get(ext, 'application/octet-stream')

    def _is_valid_mime_type(self, mime_type: str) -> bool:
        """Validate MIME type against allowed types."""
        allowed_mimes = {
            'text/plain': ['txt'],
            'application/pdf': ['pdf'],
            'application/msword': ['doc'],
            (
                'application/vnd.openxmlformats-officedocument.'
                'wordprocessingml.document'
            ): ['docx'],
        }
        return any(
            ext in settings.ALLOWED_EXTENSIONS
            for ext in allowed_mimes.get(mime_type, [])
        )

    async def _verify_file_hash(
        self,
        file_path: Path,
        expected_hash: str
    ) -> bool:
        """Verify file content hash asynchronously."""
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()
            return hashlib.sha256(content).hexdigest() == expected_hash

    async def _remove_file(self, file_path: Path) -> None:
        """Remove a file asynchronously."""
        try:
            os.remove(file_path)  # os.remove is fast enough to not need async
        except (OSError, IOError):
            pass  # Ignore file removal errors during cleanup
